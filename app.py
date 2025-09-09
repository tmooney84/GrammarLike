from flask import Flask, request, jsonify, send_file, render_template
from celery import Celery
import os
import nltk, subprocess, textwrap

nltk.download("punkt", quiet=True)
app = Flask(__name__)

app.config['CELERY_BROKER_URL'] = 'redis://localhost:6379/0'
app.config['CELERY_RESULT_BACKEND'] = 'redis://localhost:6379/0'
celery = Celery(app.name, broker=app.config['CELERY_BROKER_URL'])
celery.conf.update(app.config)

UPLOAD_FOLDER = "uploads"
PROCESSED_FOLDER = "processed"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)


def split_into_chunks(text, max_chars=3000):
    sentences = nltk.sent_tokenize(text)
    chunks, current = [], ""
    for s in sentences:
        if len(current) + len(s) + 1 > max_chars:
            chunks.append(current.strip())
            current = s
        else:
            current += " " + s
    if current:
        chunks.append(current.strip())
    return chunks


@celery.task
def process_docx(file_path, out_path, model="mistral", ctx_size=8000):
    # Extract text from uploaded docx (stub: replace with real extraction if needed)
    import docx, requests, json

    def run_ollama(prompt, model="mistral", ctx_size=8000):
        try:
            resp = requests.post(
                "http://localhost:11434/api/generate",
                json={"model": model, "prompt": prompt, "options": {"num_ctx": ctx_size}},
                stream=True,  # Enable streaming
            )
            resp.raise_for_status()  # Raise an error for bad responses

            # Initialize an empty list to accumulate responses
            full_response = []

            # Process the streamed response
            for line in resp.iter_lines():
                if line:  # Ensure the line is not empty
                    json_line = line.decode('utf-8')  # Decode the line
                    try:
                        data = json.loads(json_line)  # Parse the JSON
                        if data.get("done"):
                            break  # Stop if the response indicates completion
                        full_response.append(data.get("response", ""))  # Accumulate the response
                    except json.JSONDecodeError:
                        print("Failed to decode JSON:", json_line)  # Log any decoding errors

            # Join the accumulated responses into a single string
            return ''.join(full_response).strip()

        except requests.exceptions.HTTPError as http_err:
            print(f"HTTP error occurred: {http_err}")  # Log HTTP errors
        except Exception as err:
            print(f"An error occurred: {err}")  # Log any other errors

        return "Error: Invalid response from the server."

    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    # Chunking
    chunks = split_into_chunks(text, max_chars=(ctx_size - 500) * 4)

    # Proofread each chunk with Ollama
    corrected = []
    for chunk in chunks:
        prompt = textwrap.dedent(f"""
        You are an expert editor. Proofread the following document. Make only necessary corrections.  
        Rules:  
                                    1. Do not add extra explanations outside the text.  
                                    2. Preserve the original text structure.  
                                    3. Only correct errors or improve clarity.  
                                    Text to proofread: {chunk}
                                    """)
        result = run_ollama(prompt)
        corrected.append(result.strip())

    # Reassemble and save processed docx
    new_doc = docx.Document()
    for para in "\n\n---\n\n".join(corrected).splitlines():
        new_doc.add_paragraph(para)
    new_doc.save(out_path)

    return out_path


@app.route("/")
def index():
    ###may need some work
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    file = request.files["file"]
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    ### maybe parse the so instead of file.docx.zip just file.zip
    output_path = os.path.join(PROCESSED_FOLDER, f"processed_{file.filename}")
    file.save(file_path)
    task = process_docx.delay(file_path, output_path)
    return jsonify({"task_id": task.id})


@app.route("/status/<task_id>")
def task_status(task_id):
    task = celery.AsyncResult(task_id)
    if task.state == "SUCCESS":
        return jsonify({"status": "done", "download_url": f"/download/{task.id}"})
    return jsonify({"status": task.state})


@app.route("/download/<task_id>")
def download(task_id):
    task = celery.AsyncResult(task_id)
    if task.state == "SUCCESS":
        return send_file(task.result, as_attachment=True)
    return "Not ready", 404


###Maybe have this trigger the delete file after the download is finished...

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=3000, debug=True)
